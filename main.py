from groq import Groq
import streamlit as st
import pdfplumber
import docx
from docx import Document
from io import BytesIO
import re
import streamlit.components.v1 as components

# === Resume cleanup filter ===
def clean_output(text):
    text = re.sub(r'(\w)\1{2,}', r'\1', text)  # Remove excessive letter repetition
    blacklist = [
        r"god at", r"guru", r"ninja", r"wizard", r"rockstar",
        r"tony stark", r"batman", r"avenger", r"superhero"
    ]
    for pattern in blacklist:
        text = re.sub(pattern, "", text, flags=re.IGNORECASE)
    return text.strip()

# === Groq API setup ===
client = Groq(api_key="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")

# === Streamlit App Setup ===
st.set_page_config(page_title="Resume Weaponizer v2", layout="centered")

st.title("🧠 Resume Weaponizer v2")
st.write("Tailor your resume with Groq-powered precision. Truthful, specific, export-ready.")

# === Upload ===
resume_file = st.file_uploader("📄 Upload your resume (PDF or DOCX)", type=["pdf", "docx"])
job_desc = st.text_area("📌 Paste job description")

# === Resume Extraction ===
def extract_text(file):
    if file.name.endswith(".pdf"):
        with pdfplumber.open(file) as pdf:
            return "\n".join(page.extract_text() for page in pdf.pages if page.extract_text()), "pdf"
    elif file.name.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join(para.text for para in doc.paragraphs), "docx"
    return "", ""

# === DOCX Generator ===
def generate_docx(content):
    doc = Document()
    for line in content.split('\n'):
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# === System Prompt ===
system_prompt = """
You are a professional resume writer.

Your job is to rewrite a candidate's resume based on a job description. Keep it persuasive, confident, and tailored — but DO NOT exaggerate unrealistically or invent fictional titles or phrases.

🚫 Absolutely DO NOT use:
- "God at" anything
- Superhero or movie references
- Unrealistic job titles (e.g., "Digital Sorcerer", "Marketing Ninja")
- Over-the-top claims (e.g., "Revolutionized the internet", "Tony Stark of IT")

✅ Instead, focus on:
- Real achievements and experience
- Strong action verbs (e.g., led, implemented, optimized, designed)
- Matching skills to job requirements
- Professional tone and clarity
- Believable, grounded confidence

Use bullet points for experience, and clear formatting. Return ONLY the rewritten resume.
"""

# === Resume Enhancement ===
if st.button("🚀 Enhance Resume"):
    if resume_file and job_desc:
        resume_text, ext = extract_text(resume_file)
        full_prompt = f"{system_prompt}\n\nResume:\n{resume_text}\n\nJob Description:\n{job_desc}"

        with st.spinner("Rewriting resume with Groq..."):
            completion = client.chat.completions.create(
                model="llama3-70b-8192",
                messages=[{"role": "user", "content": full_prompt}]
            )
            enhanced = clean_output(completion.choices[0].message.content)

        st.success("🔥 Resume upgraded!")

        st.text_area("📄 Optimized Resume", enhanced, height=400, key="resume_box")

        # Copy button using JS
        components.html(f"""
            <textarea id="resumeText" style="width:100%; height:0px; opacity:0;">{enhanced}</textarea>
            <button onclick="navigator.clipboard.writeText(document.getElementById('resumeText').value)"
            style="padding: 0.5em 1em; background-color: #256fb3; color: white; border: none; border-radius: 5px; margin-top: 10px;">
                📋 Copy to Clipboard
            </button>
        """, height=60)

        # Always allow DOCX export
        docx_data = generate_docx(enhanced)
        st.download_button("⬇️ Export as DOCX", data=docx_data, file_name="enhanced_resume.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    else:
        st.warning("Upload a resume and paste the job description.")
